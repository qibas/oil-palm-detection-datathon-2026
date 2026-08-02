"""Bab 3 "Solusi Usulan" (ringkas, tingkat konsep) -> section_methodology.docx

REVISI 2026-07-23 (ketiga): dipadatkan atas permintaan Brian. Paper ini masih
paper KONSEP, jadi Bab 3 menjelaskan RANCANGAN dan temuan pokoknya, bukan
seluruh detail implementasi. Detail lengkapnya tetap hidup di 00_HASIL.md.

Yang dipertahankan utuh karena sudah disetujui:
  - paragraf pembuka Bab 3
  - paragraf pertama 3.1 (data Layer 1)
  - paragraf pertama 3.2 (rekonstruksi geometri)

Aturan gaya (register acuan "Laporan_Timnya Olip - Lisa Olivia"):
  1. TIDAK ADA EM DASH. Pemisah = koma, titik, kurung, atau konektor.
  2. Kalimat deklaratif datar. Tanpa cetak tebal di badan teks.
  3. Istilah Inggris dibiarkan apa adanya dan dicetak miring.

Sitasi ditulis sebagai penanda nama dalam kurung siku, mis. [Bonneau], supaya
Brian tinggal mengganti dengan nomor daftar pustaka. Padanannya ada di
`REFERENSI.md` bagian "Padanan penanda sitasi".
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
# Bisa dialihkan lewat env var, berguna ketika section_methodology.docx sedang
# dibuka di Word (Word mengunci berkasnya, jadi penulisan langsung gagal).
DOCX = os.environ.get("METHODOLOGY_DOCX", os.path.join(HERE, "section_methodology.docx"))

FONT, BODY_PT, TABLE_PT, CAP_PT = "Times New Roman", 10.5, 8.0, 8.5
INK, MUTED = RGBColor(0, 0, 0), RGBColor(0x40, 0x40, 0x40)

BOLD, ITAL = "*", "~"      # *tebal* dan ~miring~
EM_DASH = "—"              # dilarang muncul di ISI; dijaga assert di main()


def set_font(run, size=BODY_PT, color=INK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def runs(par, text, size=BODY_PT, color=INK, base_bold=False):
    bold, ital, buf = base_bold, False, ""

    def flush():
        nonlocal buf
        if buf:
            r = par.add_run(buf)
            r.bold, r.italic = bold, ital
            set_font(r, size, color)
            buf = ""

    for ch in text:
        if ch == BOLD:
            flush()
            bold = not bold
        elif ch == ITAL:
            flush()
            ital = not ital
        else:
            buf += ch
    flush()
    return par


def para(doc, text, space_before=0, space_after=5.0, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(space_before), Pt(space_after)
    pf.line_spacing = 1.0
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    return runs(p, text)


def heading(doc, text, size, space_before=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(space_before), Pt(2.5)
    pf.line_spacing = 1.0
    pf.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    set_font(r, size)
    return p


# --------------------------------------------------------------------------
# ISI
# --------------------------------------------------------------------------
OPENING = (
    "Pipeline disusun sebagai enam tahap berurutan, dari piksel sampai peluang, dan dirangkum "
    "pada Gambar 1. Lapisan pertama bekerja pada citra UAV dan menghasilkan inventaris tajuk. "
    "Lapisan kedua merekonstruksi tata letak kebun yang sebenarnya dari inventaris tersebut. Dua "
    "tahap terakhir menjalankan dinamika penularan di atas tata letak nyata itu dan "
    "menerjemahkan keluaran model menjadi persentase risiko yang terkalibrasi."
)

FIGNOTE = "Gambar 1: Arsitektur pipeline. [sisipkan gambar di sini]"

L1 = [
    ("Layer 1 memakai dataset ~Oil Palm Health Detection~ yang terdiri atas tiga ortomosaik UAV "
     "RGB bersudut nadir, 2.303 ubin berukuran 1.024 piksel, dengan GSD sekitar 8,7 cm/px. Label "
     "yang tersedia berupa dua kelas, yaitu ~Healthy~ dan ~Unhealthy~, yang merepresentasikan "
     "kesehatan tajuk."),
    ("Tiga tahap berjalan berurutan di atas ubin citra tersebut. Deteksi tajuk memakai YOLO11 "
     "[11] dan menghasilkan kotak pembatas beserta koordinat piksel pusat tajuk; koordinat "
     "inilah, bukan sekadar hitungan jumlah pohon, yang menjadi tulang punggung rekonstruksi "
     "geometri kebun pada Subbab 3.2. Estimasi luas tajuk memisahkan piksel tajuk dari piksel "
     "tanah memakai indeks ~Excess Green~ [14] yang diambangkan dengan metode Otsu "
     "[12]; karena data hanya menyediakan kotak pembatas tanpa ~mask~ kebenaran-dasar, tahap "
     "ini diperlakukan sebagai demonstrasi kualitatif dan tidak ada klaim IoU yang diajukan. "
     "Penilaian kesehatan tajuk mengklasifikasikan statistik warna dan tekstur kanal RGB di "
     "dalam tiap kotak dengan LightGBM [13], dengan fitur yang dipilih atas alasan agronomis, "
     "sebab klorosis termanifestasi sebagai pergeseran kehijauan."),
    ("Protokol evaluasinya ditentukan oleh keterbatasan data. Hanya tersedia tiga ortomosaik dan "
     "ubinnya saling bertindih rapat, sehingga pembagian acak dipastikan bocor dan seluruh "
     "evaluasi memakai validasi silang blok ~leave-one-ortho-out~. Metriknya mAP untuk deteksi "
     "serta PR-AUC dan ROC-AUC untuk kesehatan, sedangkan akurasi tidak dilaporkan karena pada "
     "66 positif dari 5.077 pohon sebuah pengklasifikasi yang selalu menjawab ~Healthy~ sudah "
     "melampaui 98% akurasi tanpa kegunaan apa pun."),
]

L2 = [
    ("Layer 2 dimulai dengan merekonstruksi tata letak kebun dari inventaris tajuk. Nama berkas "
     "tiap ubin menyimpan offset piksel absolutnya di ortomosaik induk, sehingga penjumlahan "
     "offset dengan pusat kotak pembatas memberi koordinat global pada setiap tajuk. Deteksi "
     "ganda antar-ubin digabungkan memakai aturan jarak, menyusutkan 151.060 kotak anotasi "
     "menjadi 5.077 pohon unik. Skala metrik diperoleh dengan memadankan jarak tetangga terdekat "
     "101 sampai 106 piksel terhadap jarak tanam lazim 9 m."),
    ("Dinamika penularan diestimasi pada panel lapangan Eg9PP [2] yang berisi 1.200 sawit, 14 "
     "famili, 2 ~parcel~, dan 45 sensus sepanjang 25 tahun, dengan gejala ~Ganoderma~ "
     "terverifikasi lapangan per pohon dan tanpa citra sama sekali. Setelah koreksi skala sumbu, "
     "keenam tetangga terdekat jatuh tepat pada jarak 1,000, menandakan pola tanam segitiga sama "
     "sisi. Di atasnya dibangun graf kedekatan dengan sisi antar pohon berjarak paling jauh 1,5 "
     "kali jarak tanam, menghasilkan 3.354 sisi dengan derajat rata-rata 5,59. Relasinya disebut "
     "kedekatan dan bukan kontak akar, sesuai alasan yang diuraikan pada Subbab 2.1."),
    ("Kedua sumber data tidak digabungkan, karena tidak ada himpunan data yang memuat citra UAV "
     "sekaligus riwayat penyebaran per pohon. Yang diukur adalah kesesuaian bentuknya. Derajat "
     "rata-rata graf pada radius yang sama bernilai 5,62 pada geometri Layer 1 dan 5,74 pada "
     "Eg9PP, berselisih 2%. Ini pemeriksaan kewarasan dan bukan bukti keteralihan, sebab pada "
     "kisi segitiga cangkang tetangga pertama berisi enam pohon pada jarak 1,0 sementara cangkang "
     "berikutnya baru muncul pada 1,73, sehingga radius 1,5 memaksa derajat mendekati enam untuk "
     "kebun mana pun yang ditanam segitiga [17]."),
    ("Tugas peramalan dirumuskan sebagai berikut. Pohon yang masih asimptomatik pada sensus ~t~ "
     "masuk ke dalam ~risk set~, dengan target apakah ia menjadi simptomatik atau mati dalam ~h~ "
     "sensus berikutnya, untuk ~h~ = 1 sampai 4. Pohon yang sudah bergejala, mati, atau tersensor "
     "dikeluarkan, dan pohon tersensor tidak dianggap sehat. Setiap model membaca jendela tiga "
     "sensus terakhir, dan genotipe wajib menjadi kovariat pada semua model termasuk garis dasar "
     "karena famili sekerabat ditanam berdekatan."),
    ("Tiga model diadu pada tugas identik, yaitu MLP tanpa graf dengan 3.713 parameter, STGNN "
     "yang memakai difusi tetangga dan GRU [20] dengan 8.875 parameter, serta STGNN+SI(D) yang "
     "menambahkan kepala epidemiologi terlatih sebesar tiga parameter saja, menjadi 8.878. "
     "MLP menerima seluruh "
     "informasi non-relasional tetapi tidak mengetahui siapa bertetangga dengan siapa, sehingga "
     "selisihnya terhadap model bergraf mengisolasi kontribusi struktur. Kompartemen laten ~E~ "
     "tidak teramati di lapangan, sehingga kepala SEIR diturunkan menjadi SI(D). Keluaran "
     "ketiganya dinilai sebagai peringkat dengan AUC-PR, sedangkan kalibrasi skor menjadi "
     "persentase belum dikerjakan sehingga belum diklaim."),
    ("Selisih STGNN dikurangi MLP dipecah dengan menukar tampilan graf sambil menahan arsitektur "
     "tetap sama, menjadi temporal (tanpa-graf − MLP), prevalensi (acak − tanpa-graf), dan "
     "struktur (asli − acak), dengan tampilan acak mempertahankan derajat tiap pohon tetapi "
     "menghancurkan strukturnya. Karena tampilan acak global menghubungkan pohon berjarak median "
     "13,2 jarak tanam, ditambahkan tangga kontrol acak-lokal yang membatasi setiap sisi baru "
     "pada 6 dan 3 jarak tanam agar lokalitas tidak tertukar dengan ketepatan peta. Evaluasinya "
     "memakai ~leave-one-parcel-out~ dua lipatan dengan nol sisi melintas ~parcel~, dijalankan "
     "pada 20 ~seed~ sehingga menghasilkan 40 pasangan, dan aturan keputusannya ditetapkan "
     "sebelum hasil dilihat, yaitu selisih yang berada di dalam satu simpangan baku dinyatakan "
     "tidak konklusif."),
]

RES_LEAD = (
    "Pada tahap konsep ini yang dilaporkan hanyalah bukti dari Layer 2, sebab di situlah premis "
    "pendekatan berbasis graf dipertaruhkan. Hasil Layer 1 tidak diuraikan karena deteksi tajuk "
    "dari citra UAV sudah merupakan persoalan yang mapan, sedangkan pertanyaan yang belum "
    "terjawab adalah apakah peta tetangga benar-benar membantu meramalkan penularan. Rincian "
    "dekomposisinya per horizon disajikan pada Lampiran A."
)

RES_AFTER = [
    ("Satu-satunya komponen yang bertahan adalah struktur, dan kekuatannya justru menguat seiring "
     "bertambahnya horizon, yaitu dari +0,0044 AUC-PR pada ~h~ = 1 dengan 30 dari 40 pasangan "
     "menjadi +0,0165 pada ~h~ = 4 dengan 39 dari 40 pasangan. Efek tersebut bertahan terhadap "
     "tiga kontrol, yaitu tampilan acak yang mempertahankan derajat setiap pohon, ~null~ "
     "permutasi yang diacak di dalam famili, serta tangga lokalitas yang menyisakan 78 sampai "
     "85% efeknya ketika graf pembanding dipaksa tetap lokal."),
    ("Sisanya tidak mendukung. Komponen temporal dan prevalensi tidak konklusif menurut aturan "
     "yang ditetapkan sebelum hasil dilihat, dan yang temporal bahkan meluruh mendekati nol pada "
     "horizon terjauh. Sebabnya struktural, yaitu Eg9PP tidak memiliki citra sehingga riwayat "
     "pohon selagi asimptomatik praktis tidak membawa informasi. Penambahan kepala epidemiologi "
     "terlatih justru merugikan pada keempat horizon dan kerugiannya membesar dari −0,0115 "
     "menjadi −0,0426, sementara empat upaya menaikkan mutunya seluruhnya gagal. Bukti penularan "
     "spasial pada tingkat asosiasi juga melemah, karena risiko relatif gabungan sebesar 4,47 "
     "kali turun menjadi 1,65 kali setelah stratifikasi Mantel-Haenszel per sensus [19]."),
    ("Keluaran sistem bagi pengguna ada dua, sebagaimana ditandai pada Gambar 1. Yang pertama "
     "adalah daftar intervensi berperingkat, berisi ID pohon, koordinatnya, dan skor risikonya, "
     "yang dipotong pada ~k~ teratas sesuai jumlah pohon yang sanggup diperiksa dalam satu "
     "siklus. Yang kedua adalah peta risiko tingkat blok, yaitu skor yang ditumpangkan pada kisi "
     "tanam hasil Tahap 4 sehingga sensus lapangan berikutnya dapat diarahkan ke blok yang "
     "paling berisiko. Koordinat itulah alasan Layer 1 tetap diperlukan, sebab skor risiko tanpa "
     "lokasi tidak dapat menggerakkan tim ke lapangan. Pada tahap konsep ini keduanya belum "
     "dihasilkan, karena keluaran model masih berupa peringkat dan kalibrasi skor menjadi "
     "persentase belum dikerjakan."),
]

# -- Lampiran A ------------------------------------------------------------
# Tabel dekomposisi dipindahkan ke sini (keputusan Brian 2026-07-24): delapan
# dari enam belas selnya melaporkan nihil, jadi 3.3 cukup membawa dua lintasan
# pokoknya dalam prosa. Hitungan tanda titik-ujungnya tetap disebut di 3.3.
APX_LEAD = (
    "Nilai pada tabel adalah selisih AUC-PR terhadap pembanding masing-masing komponen, dirata-"
    "ratakan pada 20 ~seed~ dan 2 lipatan ~leave-one-parcel-out~. Angka dalam kurung adalah "
    "hitungan tanda, yaitu banyaknya pasangan dari 40 yang selisihnya positif. Kolom kepala SI(D) "
    "adalah selisih STGNN+SI(D) dikurangi STGNN, sehingga nilai negatif berarti kepala "
    "epidemiologi terlatih memperburuk peringkat."
)

APX_HEAD = ["h", "temporal", "prevalensi", "struktur", "kepala SI(D)"]
APX_ROWS = [
    ["1", "+0,0044 (34/40)", "+0,0026 (32/40)", "+0,0044 (30/40)", "−0,0115"],
    ["2", "+0,0034 (27/40)", "+0,0029 (33/40)", "+0,0098 (37/40)", "−0,0230"],
    ["3", "+0,0021 (21/40)", "+0,0021 (25/40)", "+0,0151 (39/40)", "−0,0293"],
    ["4", "+0,0011 (19/40)", "+0,0018 (26/40)", "+0,0165 (39/40)", "−0,0426"],
]
APX_CAP = (
    "Tabel A1: Dekomposisi selisih STGNN dikurangi MLP menjadi komponen temporal, prevalensi, "
    "dan struktur, beserta kontribusi kepala SI(D), pada keempat horizon."
)

ALL_TEXT = ([OPENING, FIGNOTE, RES_LEAD, APX_LEAD, APX_CAP] + L1 + L2 + RES_AFTER
            + [c for row in APX_ROWS for c in row] + APX_HEAD)


def main():
    bad = [t[:60] for t in ALL_TEXT if EM_DASH in t]
    assert not bad, "EM DASH masih ada di: %s" % bad

    body = [OPENING, FIGNOTE, RES_LEAD, APX_LEAD, APX_CAP] + L1 + L2 + RES_AFTER
    fat = [t[:60] for t in body if BOLD in t]
    assert not fat, "CETAK TEBAL di badan teks: %s" % fat

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(1.9)
    sec.left_margin = sec.right_margin = Cm(2.0)

    st = doc.styles["Normal"]
    st.font.name, st.font.size = FONT, Pt(BODY_PT)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(5.0)
    st.paragraph_format.line_spacing = 1.0

    heading(doc, "3. Solusi Usulan", 13, space_before=0)
    para(doc, OPENING)
    p = para(doc, FIGNOTE, space_before=1, space_after=4,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        set_font(r, CAP_PT, MUTED)
        r.italic = True

    heading(doc, "3.1  Layer 1: Persepsi dari Citra UAV", 11)
    for t in L1:
        para(doc, t)

    heading(doc, "3.2  Layer 2: Data Lapangan dan Konstruksi Graf", 11)
    for t in L2:
        para(doc, t)

    heading(doc, "3.3  Hasil dan Keluaran Sistem", 11)
    para(doc, RES_LEAD)
    for t in RES_AFTER:
        para(doc, t)

    # -- Lampiran A pada halaman tersendiri --------------------------------
    doc.add_page_break()
    heading(doc, "Lampiran A. Dekomposisi Selisih STGNN dikurangi MLP", 11, space_before=0)
    para(doc, APX_LEAD)

    tbl = doc.add_table(rows=1 + len(APX_ROWS), cols=len(APX_HEAD))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, txt in enumerate(APX_HEAD):
        cell = tbl.cell(0, j)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        runs(cell.paragraphs[0], txt, size=TABLE_PT, base_bold=True)
    for i, row in enumerate(APX_ROWS, start=1):
        for j, txt in enumerate(row):
            cell = tbl.cell(i, j)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            runs(cell.paragraphs[0], txt, size=TABLE_PT)
    for row in tbl.rows:
        for cell in row.cells:
            pf = cell.paragraphs[0].paragraph_format
            pf.space_before, pf.space_after, pf.line_spacing = Pt(1), Pt(1), 1.0

    p = para(doc, APX_CAP, space_before=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        set_font(r, CAP_PT, MUTED)

    doc.save(DOCX)
    print("wrote", DOCX)


if __name__ == "__main__":
    main()
