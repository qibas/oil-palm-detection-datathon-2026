"""section3.tex -> section3.docx

Konverter kecil untuk SUBSET LaTeX yang benar-benar dipakai section3.tex:
\\section, \\subsection, environment figure, itemize, \\textbf, \\textit, dan math
inline sederhana. Bukan konverter LaTeX umum -- kalau .tex-nya nanti memakai
perintah baru, tambahkan penanganannya di sini (skrip akan protes, bukan diam).

Sumber kebenaran tetap section3.tex; .docx selalu dihasilkan ulang darinya.
"""
import os
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
TEX = os.path.join(HERE, "section3.tex")
DOCX = os.path.join(HERE, "section3.docx")
FIGDIR = os.path.dirname(HERE)          # figur ada di root repo

BODY_FONT, BODY_PT = "Times New Roman", 11

# Perintah math yang muncul di section3.tex -> unicode.
MATH = {
    r"\times": "×", r"\pm": "±", r"\alpha": "α",
    r"\varepsilon": "ε", r"\rightarrow": "→", r"\leq": "≤",
    r"\geq": "≥", r"\cdot": "·", r"\approx": "≈",
}

# Word tidak punya rujukan silang otomatis di sini, jadi nomornya dipatok.
# Tambahkan entri kalau .tex nanti punya figur/tabel baru.
LABELS = {"fig:pipeline": "1", "tab:data": "1"}
_unknown_refs = set()


def read_group(s, i):
    """s[i] == '{'; kembalikan (isi, indeks setelah '}') dengan brace matching."""
    assert s[i] == "{", f"expected {{ at {i}"
    depth, j = 0, i
    while j < len(s):
        if s[j] == "{":
            depth += 1
        elif s[j] == "}":
            depth -= 1
            if depth == 0:
                return s[i + 1:j], j + 1
        j += 1
    raise ValueError(f"brace tidak tertutup mulai indeks {i}")


def parse_inline(s, bold=False, italic=False):
    """-> [(teks, bold, italic)]. Rekursif, jadi \\textbf{... \\textit{x}} benar."""
    out, buf, i = [], "", 0

    def flush():
        nonlocal buf
        if buf:
            out.append((buf, bold, italic))
            buf = ""

    while i < len(s):
        if s.startswith(r"\textbf{", i):
            flush()
            inner, i = read_group(s, i + len(r"\textbf"))
            out += parse_inline(inner, True, italic)
        elif s.startswith(r"\textit{", i):
            flush()
            inner, i = read_group(s, i + len(r"\textit"))
            out += parse_inline(inner, bold, True)
        elif s[i] == "$":
            flush()
            j = s.index("$", i + 1)
            math = s[i + 1:j]
            for cmd, ch in MATH.items():
                math = math.replace(cmd, ch)
            math = re.sub(r"\s+", " ", math).strip()
            math = math.replace("-", "−")        # minus, bukan hyphen
            out.append((math, bold, True))       # variabel math dicetak miring
            i = j + 1
        else:
            buf += s[i]
            i += 1
    flush()
    return out


def clean(s):
    """Normalisasi teks LaTeX -> teks biasa (dijalankan SEBELUM parse_inline)."""
    def _ref(m):
        key = m.group(1)
        if key not in LABELS:
            _unknown_refs.add(key)
        return LABELS.get(key, "??")

    s = re.sub(r"\\ref\{([^}]+)\}", _ref, s)
    s = s.replace("{,}", ",").replace("~", " ")
    s = s.replace("\\%", "%").replace("\\&", "&").replace("\\_", "_")
    s = s.replace("``", "“").replace("''", "”")
    s = s.replace("---", "—").replace("--", "–")
    return re.sub(r"[ \t]+", " ", s).strip()


def add_runs(par, text):
    for chunk, b, it in parse_inline(clean(text)):
        r = par.add_run(chunk)
        r.bold, r.italic = b, it
    return par


def caption(doc, prefix, tex_text):
    """Paragraf keterangan (gambar/tabel): kecil, abu-abu, rata tengah."""
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = cap.add_run(prefix)
    lead.bold = True
    lead.font.size = Pt(9)
    for chunk, b, it in parse_inline(clean(tex_text)):
        r = cap.add_run(chunk)
        r.bold, r.italic, r.font.size = b, it, Pt(9)
        r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return cap


def add_table(doc, flat):
    """environment table -> keterangan di atas + tabel Word ber-grid.

    Hanya menangani bentuk yang dipakai section3.tex: satu \\begin{tabular},
    baris dipisah '\\\\', sel dipisah '&', '\\hline' diabaikan.
    """
    cap_txt, _ = read_group(flat, flat.index(r"\caption") + len(r"\caption"))
    caption(doc, "Tabel 1. ", cap_txt)

    m = re.search(r"\\begin\{tabular\}\{[^}]*\}(.*?)\\end\{tabular\}", flat, re.S)
    if not m:
        sys.exit("FATAL: environment table tanpa tabular yang dapat dibaca")
    body = m.group(1).replace(r"\hline", " ")

    rows = []
    for raw in body.split(r"\\"):
        if not raw.strip():
            continue
        rows.append([c.strip() for c in re.split(r"(?<!\\)&", raw)])
    if not rows:
        sys.exit("FATAL: tabular kosong")

    ncol = max(len(r) for r in rows)
    if any(len(r) != ncol for r in rows):
        sys.exit(f"FATAL: jumlah kolom tidak seragam ({[len(r) for r in rows]})")

    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cellsrc in enumerate(row):
            cell = t.cell(i, j)
            cell.text = ""
            par = cell.paragraphs[0]
            for chunk, b, it in parse_inline(clean(cellsrc)):
                r = par.add_run(chunk)
                r.bold, r.italic, r.font.size = (b or i == 0), it, Pt(8)
    return t


def main():
    src = open(TEX, encoding="utf-8").read()
    src = "\n".join(l for l in src.splitlines() if not l.lstrip().startswith("%"))
    blocks = [b.strip() for b in re.split(r"\n\s*\n", src) if b.strip()]

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name, st.font.size = BODY_FONT, Pt(BODY_PT)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15

    unhandled = []
    for blk in blocks:
        flat = " ".join(blk.split())

        if flat.startswith(r"\section{"):
            title, _ = read_group(flat, flat.index("{"))
            add_runs(doc.add_heading(level=1), title)

        elif flat.startswith(r"\subsection{"):
            title, _ = read_group(flat, flat.index("{"))
            add_runs(doc.add_heading(level=2), title)

        elif r"\begin{figure}" in flat:
            m = re.search(r"\\includegraphics(?:\[[^\]]*\])?\{([^}]+)\}", flat)
            img = os.path.join(FIGDIR, m.group(1))
            if not os.path.exists(img):
                sys.exit(f"FATAL: figur tidak ditemukan: {img}")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(img, width=Inches(6.3))
            cap_txt, _ = read_group(flat, flat.index(r"\caption") + len(r"\caption"))
            caption(doc, "Gambar 1. ", cap_txt)

        elif r"\begin{table}" in flat:
            add_table(doc, flat)

        elif r"\begin{itemize}" in flat:
            body = flat.split(r"\begin{itemize}", 1)[1].split(r"\end{itemize}")[0]
            body = re.sub(r"\\itemsep[\d.]+pt", "", body)
            for item in [x.strip() for x in body.split(r"\item") if x.strip()]:
                add_runs(doc.add_paragraph(style="List Bullet"), item.rstrip(";."))

        # Paragraf boleh DIAWALI markup inline (mis. "\textbf{Deteksi tajuk} memakai
        # ...") -- itu teks biasa, bukan perintah struktural yang tak dikenal.
        elif flat.startswith((r"\textbf{", r"\textit{")):
            add_runs(doc.add_paragraph(), flat)

        elif flat.startswith("\\"):
            unhandled.append(flat[:70])

        else:
            add_runs(doc.add_paragraph(), flat)

    doc.save(DOCX)
    print("wrote", DOCX)
    if unhandled:
        print("PERINGATAN - blok LaTeX tidak dikenali, TIDAK ikut tertulis:")
        for u in unhandled:
            print("  ", u)
    if _unknown_refs:
        print("PERINGATAN - \\ref tanpa nomor (tertulis '??'), tambahkan ke LABELS:")
        for u in sorted(_unknown_refs):
            print("  ", u)


if __name__ == "__main__":
    main()
