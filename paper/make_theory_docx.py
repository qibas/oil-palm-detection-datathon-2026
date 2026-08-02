"""Revisi Bab 2 (Kajian Teori) + tambalan Bab 3 -> section2_kajian_teori.docx

Dibuat 2026-07-24 atas dua keputusan Brian:
  1. Istilah relasi graf memakai "kedekatan", bukan "kontak akar".
  2. Ruang mepet, jadi subbab 2.3 yang sempat direncanakan TIDAK dibuat berdiri
     sendiri. Isinya dilebur: penjahitan ubin dan pseudo-replikasi masuk 2.2,
     pola tanam dan alasan radius 1,5 masuk 2.3 (STGNN), sehingga jumlah subbab
     tetap empat seperti naskah aslinya.

Berkas ini juga membawa dua tambalan untuk Bab 3, karena naskah versi Brian
kehilangan Tahap 4 dan uji antarmuka: pembukanya menjanjikan enam tahap tetapi
hanya lima yang diuraikan, sedangkan 3.1 dan 3.3 sama-sama merujuk Tahap 4.

Sitasi baru DITAMBAHKAN sebagai [15] sampai [20] agar penomoran [1] sampai [14]
yang sudah ada tidak perlu digeser. Kalau template menuntut urutan kemunculan
yang ketat, seluruh naskah harus dinomori ulang; itu pekerjaan terpisah.

Aturan gaya sama dengan make_methodology_docx.py: tanpa em dash, tanpa cetak
tebal di badan teks, istilah Inggris dicetak miring.
"""
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.environ.get("THEORY_DOCX", os.path.join(HERE, "section2_kajian_teori.docx"))

FONT, BODY_PT, CAP_PT = "Times New Roman", 10.5, 8.5
INK, MUTED = RGBColor(0, 0, 0), RGBColor(0x40, 0x40, 0x40)
BOLD, ITAL, EM_DASH = "*", "~", "—"


def set_font(run, size=BODY_PT, color=INK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)


def runs(par, text, size=BODY_PT, color=INK):
    bold, ital, buf = False, False, ""

    def flush():
        nonlocal buf
        if buf:
            r = par.add_run(buf)
            r.bold, r.italic = bold, ital
            set_font(r, size, color)
            buf = ""

    for ch in text:
        if ch == BOLD:
            flush()
            bold = not bold
        elif ch == ITAL:
            flush()
            ital = not ital
        else:
            buf += ch
    flush()
    return par


def para(doc, text, space_before=0, space_after=5.0, size=BODY_PT, color=INK, align=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(space_before), Pt(space_after)
    pf.line_spacing = 1.0
    p.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    return runs(p, text, size=size, color=color)


def heading(doc, text, size, space_before=8):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before, pf.space_after = Pt(space_before), Pt(2.5)
    pf.line_spacing = 1.0
    pf.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    set_font(r, size)
    return p


# --------------------------------------------------------------------------
# BAB 2 — KAJIAN TEORI (revisi)
# --------------------------------------------------------------------------
T21 = [
    ("BSR, penyakit busuk pangkal batang akibat ~Ganoderma boninense~, menular melalui lebih "
     "dari satu jalur. Infeksi lewat kontak langsung antar sistem perakaran pohon yang "
     "berdekatan terdokumentasi secara mikroskopis, mulai dari penetrasi akar sampai invasi "
     "pangkal batang [15]. Namun basidiospora juga terbukti berperan sebagai sumber inokulum, "
     "dan pohon bergejala yang bertetangga kerap membawa isolat yang berbeda secara genetik, "
     "sehingga penularan vegetatif lewat akar tidak dapat diklaim sebagai jalur tunggal [16]. "
     "Kedua temuan itu tidak saling meniadakan. Keduanya sama-sama menempatkan kedekatan "
     "spasial antar pohon sebagai faktor penentu pola penyebaran, dan hanya yang pertama yang "
     "menamainya kontak akar. Penelitian ini karena itu merepresentasikan kebun sebagai graf "
     "kedekatan yang berlaku sebagai proksi kontak, yaitu tiap pohon sebagai simpul dan jarak "
     "di bawah ambang tertentu sebagai sisi, tanpa mengklaim satu mekanisme tunggal. Kesimpulan "
     "yang ditarik tetap sahih pada kedua penjelasan."),
    ("Sifat kedua yang menentukan rumusan masalah adalah fase asimptomatik yang panjang. "
     "Jaringan pengangkut air di pangkal batang melapuk secara bertahap, sementara tajuk masih "
     "tampak normal selama berbulan-bulan hingga bertahun-tahun [1], [3]. Gejala eksternal baru "
     "terlihat setelah kerusakan sudah lanjut, pada titik pohon praktis tidak dapat "
     "diselamatkan [1]. Ini menjustifikasi rumusan masalah, yaitu bukan mendeteksi pohon yang "
     "sudah sakit, melainkan meramalkan pohon sehat mana yang berisiko tertular berikutnya."),
    ("Karakteristik ketiga menyangkut cara data lapangan longitudinal diperlakukan. Tisné dkk. "
     "[2] menunjukkan bahwa data kesehatan pohon sawit yang diamati berkala selama puluhan "
     "tahun paling tepat dianalisis sebagai data ketahanan, karena sebagian pohon belum "
     "bergejala saat pengamatan berakhir, bukan karena pasti sehat. Fenomena ini disebut "
     "penyensoran kanan: pohon yang belum bergejala harus dikeluarkan dari himpunan risiko, "
     "bukan dilabeli sehat."),
    ("Populasi yang diamati pada studi semacam ini umumnya bersifat multiparental, sengaja "
     "dirancang memvariasikan ketahanan genetik [2]. Kekerabatan genetik karena itu berpotensi "
     "menjadi perancu, sebab famili yang lebih rentan mungkin ditanam berdekatan sehingga "
     "korelasi spasial gejala bisa mencerminkan kesamaan genetik dan bukan semata penularan. "
     "Genotipe perlu diperlakukan sebagai kovariat yang dikendalikan pada semua model, termasuk "
     "garis dasar."),
]

T22 = [
    ("Tahap pertama adalah menemukan dan menandai posisi tiap tajuk secara presisi. Pendekatan "
     "~object detection~, khususnya keluarga arsitektur YOLO, terbukti efektif untuk tugas ini "
     "pada citra kanopi UAV. Hao dkk. [5] mendemonstrasikan model berbasis YOLO untuk "
     "mendeteksi dan menghitung pohon dari citra UAV pada skala kebun, dengan pendekatan "
     "satu-tahap yang mengutamakan efisiensi komputasi."),
    ("Tahap berikutnya menilai kondisi kesehatan tajuk. Abbas dkk. [4] merangkum bahwa spektrum "
     "RGB maupun multispektral UAV dapat menangkap perubahan warna daun terkait tingkat "
     "klorofil dan stres tanaman, dan perubahan ini dapat dipelajari model pembelajaran mesin, "
     "termasuk model berbasis ~gradient boosting~ [13]. Pendekatan berbasis fitur dengan model "
     "tabular semacam ini lebih tahan terhadap data berlabel terbatas dan ketidakseimbangan "
     "kelas ekstrem dibanding ~deep learning~ langsung pada citra. Untuk memperkirakan luas "
     "tajuk ketika anotasi hanya berupa kotak pembatas tanpa ~mask~ kebenaran-dasar, praktik "
     "yang lazim adalah memisahkan piksel tanaman dari latar memakai indeks warna. Woebbecke "
     "dkk. [14] menunjukkan indeks ~Excess Green~ memisahkan vegetasi dari tanah dan residu "
     "pada berbagai kondisi pencahayaan, dan pemisahan itu dapat diambangkan otomatis dengan "
     "metode Otsu [12] yang memilih ambang langsung dari histogram aras keabuan. Keduanya tidak "
     "memerlukan pelatihan, sehingga dapat dipakai ketika data pelatihan segmentasi tidak "
     "tersedia."),
    ("Satu langkah terakhir menentukan cara seluruh keluaran tahap penginderaan dibaca. Citra "
     "UAV lazim dipotong menjadi ubin yang saling bertindih, sehingga satu pohon fisik muncul "
     "berkali-kali pada ubin yang berbeda, dan koordinat pusat tajuk hasil deteksi perlu "
     "dikembalikan ke kerangka ortomosaik induk lalu digabungkan. Memperlakukan pengamatan "
     "berulang atas unit yang sama sebagai replikasi independen adalah pseudo-replikasi, yang "
     "menggelembungkan ukuran sampel semu dan membuat selang kepercayaan tampak jauh lebih "
     "sempit daripada yang dibenarkan data [18]. Penggabungan deteksi ganda karena itu bukan "
     "langkah kosmetik, melainkan syarat agar unit analisisnya adalah pohon dan bukan baris "
     "anotasi."),
]

T23 = [
    ("Graf kedekatan hanya bermakna bila ambang jaraknya berdasar, dan dasarnya datang dari "
     "rancangan tanam. Sawit ditanam mengikuti rancangan segitiga sama sisi, dengan kerapatan "
     "optimum 143 sampai 160 pohon per hektare atau setara jarak 8,5 sampai 9 m antar pohon "
     "[17]. Pada kisi segitiga berjarak ~d~, cangkang tetangga pertama berisi tepat enam pohon "
     "pada jarak ~d~, sedangkan cangkang kedua baru muncul pada ~d~√3 atau sekitar 1,73~d~. "
     "Ambang ~r~ = 1,5~d~ karena itu menangkap cangkang pertama secara utuh tanpa menyentuh "
     "cangkang kedua, dan derajat rata-rata sekitar enam merupakan konsekuensi rancangan tanam, "
     "bukan parameter yang disetel. Sifat yang sama pula yang membuat derajat graf keluaran "
     "tahap penginderaan dapat dibandingkan dengan derajat graf masukan tahap peramalan "
     "meskipun kedua sumber data tidak digabungkan. Kesamaan derajat merupakan syarat perlu "
     "bagi kesebangunan antarmuka, bukan syarat cukup bagi keteralihan model."),
    ("~Spatio-temporal Graph Neural Network~ (STGNN) dirancang menangkap pola temporal dan "
     "spasial sekaligus, dengan lapisan konvolusi graf yang mengagregasi informasi antar simpul "
     "bertetangga digabung lapisan pemrosesan urutan waktu seperti ~gated recurrent unit~ [20], "
     "secara konseptual mengungguli model tabular sederhana seperti MLP yang memandang tiap "
     "pohon terisolasi. Namun keunggulan ini tidak otomatis. Fujita dan Akutsu [9] menunjukkan "
     "bahwa pada peramalan epidemi antar-wilayah, penambahan struktur graf sering meningkatkan "
     "kompleksitas tanpa manfaat performa sepadan, bahkan mendorong ~overfitting~. Klaim bahwa "
     "model graf lebih unggul perlu dipisahkan dari klaim bahwa peta kedekatan itu sendiri "
     "membawa informasi tambahan."),
    ("Untuk memisahkan kedua klaim ini diperlukan kerangka pembanding yang menjaga agar selisih "
     "performa berasal dari kebenaran struktur, bukan sekadar akses ke suatu graf. Kerangka "
     "standar dalam ilmu jaringan adalah model konfigurasi dengan pengacakan berderajat-sama, "
     "dipopulerkan Maslov dan Sneppen [10]: graf diacak ulang sehingga jumlah tetangga tiap "
     "simpul tetap sama, tetapi identitas tetangganya ditukar acak. Selisih performa antara "
     "graf asli dan graf hasil pengacakan ini menjadi dasar prosedur dekomposisi, memecah "
     "keunggulan model graf menjadi komponen nilai memodelkan waktu, nilai prevalensi tetangga "
     "sakit, dan nilai struktur kedekatan yang benar dibanding acak. Satu kehati-hatian "
     "menyertainya. Pada kisi segitiga seluruh sisi asli berpanjang sama, sehingga pengacakan "
     "berderajat-sama yang dilakukan secara global akan menyambungkan pohon yang berjauhan, dan "
     "selisih asli dikurangi acak berpotensi mengukur lokalitas alih-alih ketepatan peta. "
     "Kontrolnya karena itu perlu dinaikkan bertingkat, yaitu dengan membatasi panjang sisi "
     "hasil pengacakan agar graf pembanding tetap berada di sekitar pohon aslinya."),
]

T24 = [
    ("Data berkomponen ruang memiliki sifat autokorelasi spasial: sesuatu yang berdekatan "
     "cenderung lebih mirip dibanding yang berjauhan, dirumuskan Tobler [6] sebagai Hukum "
     "Pertama Geografi. Prinsip ini berlaku pada kondisi kesehatan tanaman antar pohon "
     "bertetangga."),
    ("Konsekuensinya, validasi silang acak melanggar asumsi independensi. Ketika data dari "
     "lokasi berdekatan tersebar ke himpunan latih dan uji lewat pembagian acak, performa "
     "terukur menjadi bias optimistis. Kattenborn dkk. [7] mendemonstrasikan hal ini langsung "
     "pada segmentasi spesies pohon dari citra UAV menggunakan CNN, menunjukkan validasi silang "
     "acak dapat menggelembungkan performa hingga puluhan persen dibanding validasi silang "
     "berbasis blok spasial. Temuan ini menjustifikasi penolakan pembagian data acak dan "
     "penggantiannya dengan skema yang menahan satu blok spasial, yaitu satu ortomosaik atau "
     "satu petak, utuh pada satu sisi pembagian saja."),
    ("Tantangan kedua muncul dari ketidakseimbangan kelas ekstrem, lazim pada deteksi penyakit "
     "tanaman karena jumlah pohon sehat jauh melebihi jumlah bergejala. Pada kondisi ini akurasi "
     "menyesatkan, karena model yang selalu memprediksi sehat tetap memperoleh skor tinggi tanpa "
     "mengidentifikasi satu pun kasus positif. Saito dan Rehmsmeier [8] menunjukkan kurva "
     "~Precision-Recall~ dan luas areanya (PR-AUC) lebih informatif dibanding ROC pada dataset "
     "timpang, karena berfokus pada seberapa baik model mengidentifikasi kelas positif yang "
     "langka. PR-AUC, dibandingkan terhadap garis dasar setara prevalensi kelas positif, karena "
     "itu menjadi metrik pelaporan utama dalam penelitian ini, bukan akurasi."),
    ("Dua alat pelaporan melengkapinya. Karena tiap konfigurasi dijalankan pada banyak ~seed~, "
     "perbandingan antarmodel dilakukan berpasangan per-~seed~ dan dilaporkan bersama hitungan "
     "tanda, sebab menang tipis pada hampir seluruh ~seed~ dan menang besar pada separuh ~seed~ "
     "adalah dua pernyataan yang berbeda; selisih yang berada di dalam satu simpangan baku "
     "dinyatakan tidak konklusif menurut aturan yang ditetapkan sebelum hasil dilihat. Selain "
     "itu, efek tetangga pada data longitudinal dapat tercemar waktu kalender, karena pohon yang "
     "diamati pada sensus yang sama sama-sama menghadapi tekanan wabah yang sama. Mantel dan "
     "Haenszel [19] merumuskan estimator risiko relatif terstratifikasi yang menggabungkan bukti "
     "antar strata sambil menahan pengaruh strata itu sendiri, dan stratifikasi per sensus "
     "dipakai di sini untuk memisahkan efek tetangga dari efek waktu. Prinsip yang sama berlaku "
     "pada uji permutasi, yaitu pengacakan dilakukan di dalam famili agar kekerabatan yang "
     "ditanam berdekatan tidak terhitung sebagai penularan."),
]

# --------------------------------------------------------------------------
# TAMBALAN BAB 3
# --------------------------------------------------------------------------
PATCH_NOTE = (
    "Naskah Bab 3 versi sekarang menjanjikan enam tahap pada paragraf pembuka, tetapi hanya "
    "lima yang diuraikan. Tahap 4 tidak ada, padahal 3.1 menutup dengan merujuknya dan 3.3 "
    "menyebut kisi tanam hasil Tahap 4. Uji antarmuka juga disebut pada Batasan 1.3 tetapi "
    "angkanya tidak pernah muncul. Dua paragraf berikut menutup kedua lubang itu."
)

PATCH1_HEAD = "Tambalan 1. Sisipkan sebagai paragraf PERTAMA Subbab 3.2, sebelum paragraf Eg9PP."
PATCH1 = (
    "Rekonstruksi geometri kebun adalah tahap yang menjembatani kedua layer. Nama berkas setiap "
    "ubin menyimpan offset piksel absolutnya di dalam ortomosaik induk, sehingga menjumlahkan "
    "offset tersebut dengan pusat kotak pembatas memberi koordinat global bagi tiap tajuk, dan "
    "seluruh ubin dapat dijahit kembali menjadi satu awan titik yang merepresentasikan tata "
    "letak kebun. Deteksi ganda antar-ubin digabungkan memakai aturan jarak, dan hasilnya "
    "menentukan cara seluruh keluaran Layer 1 dibaca, yaitu 151.060 kotak anotasi ternyata hanya "
    "mewakili 5.077 pohon unik karena satu pohon muncul pada median 32 ubin. Skala metriknya "
    "diperoleh dengan memadankan jarak tetangga terdekat sebesar 101 sampai 106 piksel terhadap "
    "jarak tanam sawit yang lazim, yaitu 9 m."
)

PATCH2_HEAD = ("Tambalan 2. Sisipkan setelah kalimat pembangunan graf pada paragraf Eg9PP di "
               "Subbab 3.2. Ganti pula frasa \"graf kontak akar\" menjadi \"graf kedekatan\".")
PATCH2 = (
    "Kedua sumber data tidak digabungkan, dan yang diukur hanyalah kesebangunan antarmukanya. "
    "Derajat rata-rata graf pada radius 1,5 kali jarak tanam bernilai 5,62 pada geometri hasil "
    "Layer 1 dan 5,74 pada Eg9PP, atau berselisih 2%. Sesuai Subbab 2.3, kesepakatan sebesar itu "
    "memang diharapkan pada dua kebun yang sama-sama ditanam segitiga, sehingga ia berlaku "
    "sebagai pemeriksaan kewarasan bahwa rekonstruksi dan deduplikasi tidak merusak kerapatan "
    "lokal, bukan sebagai bukti keteralihan model. Angkanya dihitung dari kotak "
    "kebenaran-dasar dan bukan prediksi detektor, sehingga merupakan batas atas."
)

REFS_HEAD = ("Tambahan Daftar Pustaka. Dinomori [15] sampai [20] agar penomoran [1] sampai [14] "
             "yang sudah ada tidak perlu digeser.")

REFS = [
    "[15] Rees, R.W., Flood, J., Hasan, Y., Potter, U., & Cooper, R.M. (2009). Basal stem rot of "
    "oil palm (Elaeis guineensis); mode of root infection and lower stem invasion by Ganoderma "
    "boninense. Plant Pathology, 58(5), 982-989.",
    "[16] Pilotti, C.A., Gorea, E.A., & Bonneau, L. (2018). Basidiospores as sources of inoculum "
    "in the spread of Ganoderma boninense in oil palm plantations in Papua New Guinea. Plant "
    "Pathology, 67(9), 1841-1849.",
    "[17] Bonneau, X., Impens, R., & Buabeng, M. (2018). Optimum oil palm planting density in "
    "West Africa. OCL, 25(2), A201.",
    "[18] Hurlbert, S.H. (1984). Pseudoreplication and the Design of Ecological Field "
    "Experiments. Ecological Monographs, 54(2), 187-211.",
    "[19] Mantel, N., & Haenszel, W. (1959). Statistical Aspects of the Analysis of Data From "
    "Retrospective Studies of Disease. Journal of the National Cancer Institute, 22(4), 719-748.",
    "[20] Cho, K., van Merrienboer, B., Gulcehre, C., Bahdanau, D., Bougares, F., Schwenk, H., & "
    "Bengio, Y. (2014). Learning Phrase Representations using RNN Encoder-Decoder for Statistical "
    "Machine Translation. arXiv:1406.1078.",
]

BODY = T21 + T22 + T23 + T24 + [PATCH_NOTE, PATCH1, PATCH2] + REFS


def main():
    bad = [t[:60] for t in BODY if EM_DASH in t]
    assert not bad, "EM DASH masih ada di: %s" % bad
    fat = [t[:60] for t in T21 + T22 + T23 + T24 + [PATCH1, PATCH2] if BOLD in t]
    assert not fat, "CETAK TEBAL di badan teks: %s" % fat

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(1.9)
    sec.left_margin = sec.right_margin = Cm(2.0)

    st = doc.styles["Normal"]
    st.font.name, st.font.size = FONT, Pt(BODY_PT)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    st.paragraph_format.space_after = Pt(5.0)
    st.paragraph_format.line_spacing = 1.0

    heading(doc, "2. Kajian Teori", 13, space_before=0)

    for title, block in [
        ("2.1  Karakteristik Epidemiologis Basal Stem Rot (BSR)", T21),
        ("2.2  Ekstraksi Fitur Spasial dari Citra UAV", T22),
        ("2.3  Pemodelan Spasial-Temporal dan Graph Neural Networks (GNN)", T23),
        ("2.4  Tantangan Validasi pada Data Geospasial", T24),
    ]:
        heading(doc, title, 11)
        for t in block:
            para(doc, t)

    heading(doc, "Tambalan untuk Bab 3", 13)
    para(doc, PATCH_NOTE, size=CAP_PT, color=MUTED)
    for head, txt in [(PATCH1_HEAD, PATCH1), (PATCH2_HEAD, PATCH2)]:
        heading(doc, head, 10, space_before=6)
        para(doc, txt)

    heading(doc, "Tambahan Daftar Pustaka", 13)
    para(doc, REFS_HEAD, size=CAP_PT, color=MUTED)
    for r in REFS:
        p = para(doc, r, space_after=3.0, align=WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.left_indent = Cm(0.9)
        p.paragraph_format.first_line_indent = Cm(-0.9)

    doc.save(DOCX)
    print("wrote", DOCX)


if __name__ == "__main__":
    main()
